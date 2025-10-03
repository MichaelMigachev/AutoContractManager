# core/document_generator.py
from openpyxl import load_workbook

from docx import Document
from pathlib import Path
import logging
from typing import Dict, Any

# Импорты из проекта
from config.paths import OUTPUT_DIR, CONTRACT_TEMPLATE, INVOICE_TEMPLATE, INVOICE_CARD_TEMPLATE
from config.settings import (
    APP_NAME,
    COMPANY_NAME,
)
from core.utils import sanitize_filename, get_current_date, number_to_words, get_date_verbose


def replace_placeholders_in_paragraph(paragraph, data: Dict[str, Any]):
    """Заменяет {ключ} на значение в параграфе"""
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        if placeholder in paragraph.text:
            # Сохраняем форматирование каждого run
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))


def replace_placeholders_in_table(table, data: Dict[str, Any]):
    """Заменяет {ключ} на значение в таблицах"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)


def fill_template(template_path: Path, output_path: Path, data: Dict[str, Any]) -> bool:
    """
    Заполняет шаблон Word и сохраняет результат

    :param template_path: путь к .docx шаблону
    :param output_path: куда сохранить заполненный документ
    :param data: словарь с данными для подстановки
    :return: True при успехе
    """
    try:
        if not template_path.exists():
            logging.error(f"Шаблон не найден: {template_path}")
            return False

        doc = Document(template_path)

        # Замена в обычных параграфах
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data)

        # Замена в таблицах
        for table in doc.tables:
            replace_placeholders_in_table(table, data)

        # Сохранение
        doc.save(output_path)
        logging.info(f"Документ создан: {output_path}")
        return True

    except Exception as e:
        logging.error(f"Ошибка при генерации документа: {e}")
        return False


def generate_contract(client_data: Dict[str, Any]) -> bool:
    """
    Создаёт договор на основе данных клиента

    :param client_data: данные клиента из Excel
    :return: True при успехе
    """
    from core.database import get_next_contract_number

    # Данные для шаблона
    contract_num = get_next_contract_number()
    full_name = f"{client_data['Фамилия']} {client_data['Имя']} {client_data['Отчество']}"
    car_info = f"{client_data['Марка авто']} (VIN {client_data['VIN']})"
    phone, _ = client_data.get("Телефон_формат", ("", ""))  # может быть предварительно обработан
    if not phone:
        phone, _ = format_phone(client_data['Телефон'])

    context = {
        "NUM": contract_num.replace("-ИП", ""),
        "FULL_NUM": contract_num,
        "DATE": get_current_date(),
        "FIO": full_name,
        "FULL_FIO": full_name,
        "SHORT_FIO": f"{client_data['Фамилия']} {client_data['Имя'][0]}. {client_data['Отчество'][0]}.",
        "CAR": client_data['Марка авто'],
        "VIN": client_data['VIN'],
        "CAR_INFO": car_info,
        "ADDRESS": client_data['Адрес'],
        "INDEX": client_data['Индекс'],
        "PASSPORT": client_data['Паспорт (серия и номер)'],
        "ISSUED_BY": client_data['Кем выдан'],
        "ISSUE_DATE": client_data['Дата выдачи'],
        "DEP_CODE": client_data['Код подразделения'],
        "BIRTH_DATE": client_data['Дата рождения'],
        "PHONE": phone,
        "COMPANY": COMPANY_NAME,
        "APP_NAME": APP_NAME
    }

    # Имя файла
    safe_fio = sanitize_filename(full_name)
    safe_index = client_data['Индекс'] if '*' not in client_data['Индекс'] else "ИНДЕКС"
    filename = f"Договор №{contract_num} ({safe_fio})_{safe_index}.docx"
    output_path = OUTPUT_DIR / filename

    # Генерация
    success = fill_template(CONTRACT_TEMPLATE, output_path, context)
    return success


# --- Удобные функции ---

def format_phone(phone: str) -> tuple[str, str]:
    """Дублируем из utils, если нужно (или импортировать)"""
    from core.utils import format_phone as util_format
    return util_format(phone)


def generate_invoice(
    client_data: dict,
    contract_num: str,
    service_type: str,
    amount: int,
    payment_method: str = "card"
) -> bool:
    """
    Заполняет шаблон счёта с использованием {ключей}
    """
    # Выбор шаблона
    if payment_method == "card":
        template_path = INVOICE_CARD_TEMPLATE
    else:
        template_path = INVOICE_TEMPLATE

    output_filename = (f"Подписанный счет{' НА КАРТУ ' if payment_method == 'card' else ' '}№ {contract_num[:3]}-001 от"
                       f" {get_current_date()} для {client_data['Фамилия']} {client_data['Имя']} "
                       f"{client_data['Отчество']}_{client_data['Индекс']}.docx")
    output_path = OUTPUT_DIR / sanitize_filename(output_filename)

    try:
        if not template_path.exists():
            logging.error(f"Шаблон не найден: {template_path}")
            return False

        doc = Document(template_path)

        # Определяем услугу
        service_desc = "выпуску СБКТС + ЭПТС" if service_type == "sbkts" else "списанию утильсбора"

        # Формируем контекст
        current_date = get_current_date()  # "03.10.2025"
        verbose_date = get_date_verbose(current_date)  # "3 октября 2025 г."

        # Формируем контекст
        context = {
            "NUM": contract_num[:3],
            "DATE": get_current_date(),
            "VERBOSE_DATE": verbose_date,
            "FIO": f"{client_data['Фамилия']} {client_data['Имя']} {client_data['Отчество']}",
            "ADDRESS": client_data['Адрес'],
            "SERVICE": service_desc,
            "CAR": f"{client_data['Марка авто']}_vin {client_data['VIN']}",
            "AMOUNT": f"{amount:.2f}".replace('.00', ''),  # Без .00
            "AMOUNT_RUB": f"{amount} руб.",
            "AMOUNT_TEXT": number_to_words(amount // 1000).capitalize() + " тысяч",
            "CONTRACT_REF": contract_num
        }

        # --- Замена в параграфах ---
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, context)

        # --- Замена в таблицах ---
        for table in doc.tables:
            replace_placeholders_in_table(table, context)

        doc.save(output_path)
        logging.info(f"✅ Счёт создан: {output_path}")
        return True

    except Exception as e:
        logging.error(f"🔴 Ошибка генерации счёта: {e}")
        return False
